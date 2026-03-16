"""
Agente Jurídico para Criação de Documentos Padronizados
Busca templates reais no banco (ANC_SIS_JUD_DOC_MODELO), baixa o .docx
e substitui as variáveis ${...} pelos dados fornecidos.
"""

import os
import sys
import re
import gc
import html
import urllib.request
import tempfile
from copy import deepcopy
from typing import List, Dict
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .database_manager import DatabaseManager

# Configura encoding UTF-8 para Windows
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

from shared.config import DATABASE_CONFIG as DB_CONFIG
from shared.config import get_logger

logger = get_logger(__name__)


class AgenteJuridico:
    """Agente Jurídico dinâmico — busca templates do banco e preenche variáveis."""

    def __init__(self, db_config: dict):
        self.db = DatabaseManager(db_config)

    
    # ──────────────────────────────────────────────────────────────────────
    # 3. Baixar o .docx do link S3
    # ──────────────────────────────────────────────────────────────────────
    def baixar_template(self, link: str) -> str:
        """Baixa o .docx do link e retorna o caminho local do arquivo."""
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        urllib.request.urlretrieve(link, tmp.name)
        return tmp.name

    # ──────────────────────────────────────────────────────────────────────
    # 5. Substituir variáveis no .docx
    # ──────────────────────────────────────────────────────────────────────
    def _substituir_em_paragrafos(self, paragrafos, dados: Dict[str, str]) -> int:
        """Substitui variáveis ${...} nos runs dos parágrafos."""
        substituicoes = 0
        for para in paragrafos:
            # Primeiro tenta substituir run a run
            for run in para.runs:
                original = run.text
                novo = original
                for var, valor in dados.items():
                    if var in novo:
                        novo = novo.replace(var, valor)
                if novo != original:
                    run.text = novo
                    substituicoes += 1

            # Verifica se ainda restam variáveis (caso a variável esteja
            # quebrada entre múltiplos runs)
            texto_completo = para.text
            variaveis_restantes = re.findall(r'\$\{[^}]+\}', texto_completo)
            if variaveis_restantes:
                # Reconstrói o parágrafo juntando todos os runs
                texto_novo = texto_completo
                precisa_rebuild = False
                for var, valor in dados.items():
                    if var in texto_novo:
                        texto_novo = texto_novo.replace(var, valor)
                        precisa_rebuild = True

                if precisa_rebuild and para.runs:
                    # Mantém a formatação do primeiro run e limpa os demais
                    para.runs[0].text = texto_novo
                    for run in para.runs[1:]:
                        run.text = ""
                    substituicoes += 1

        return substituicoes

    def substituir_variaveis(self, doc_path: str, dados: Dict[str, str], output_path: str) -> int:
        """
        Abre o .docx, substitui todas as variáveis ${...} pelos dados
        fornecidos e salva o resultado em output_path.
        Retorna o total de substituições feitas.
        """
        doc = Document(doc_path)
        total = 0

        total += self._substituir_em_paragrafos(doc.paragraphs, dados)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total += self._substituir_em_paragrafos(cell.paragraphs, dados)

        for section in doc.sections:
            total += self._substituir_em_paragrafos(section.header.paragraphs, dados)
            total += self._substituir_em_paragrafos(section.footer.paragraphs, dados)

        doc.save(output_path)

        # Libera o documento para desbloquear o arquivo no Windows
        del doc
        gc.collect()

        return total

    # Variáveis que devem ser removidas (substituídas por vazio)
    VARIAVEIS_IGNORADAS = {"TAGS"}
    def _limpar_html(self, texto_html: str) -> str:
        """Remove tags HTML e decodifica entidades, retornando texto limpo."""
        sem_tags = re.sub(r'<[^>]+>', '', texto_html)
        return html.unescape(sem_tags).strip()

    def _processar_html(self, texto_html: str) -> list:
        """
        Processa HTML e retorna lista ordenada de elementos:
        - {"tipo": "texto", "conteudo": "texto limpo"}
        - {"tipo": "tabela", "linhas": [["col1", "col2"], ...]}
        """
        elementos = []
        partes = re.split(r'(<table[\s\S]*?</table>)', texto_html, flags=re.IGNORECASE)

        for parte in partes:
            parte = parte.strip()
            if not parte:
                continue

            if re.match(r'<table', parte, re.IGNORECASE):
                linhas = []
                rows = re.findall(r'<tr[^>]*>([\s\S]*?)</tr>', parte, re.IGNORECASE)
                for row_html in rows:
                    cells = re.findall(r'<t[dh][^>]*>([\s\S]*?)</t[dh]>', row_html, re.IGNORECASE)
                    linha = []
                    for cell_html in cells:
                        cell_text = re.sub(r'<[^>]+>', '', cell_html)
                        cell_text = html.unescape(cell_text).strip()
                        linha.append(cell_text)
                    if linha:
                        linhas.append(linha)
                if linhas:
                    elementos.append({"tipo": "tabela", "linhas": linhas})
            else:
                # Quebra por <br> / <br/> / <br /> para respeitar quebras de linha
                segmentos = re.split(r'<br\s*/?>', parte, flags=re.IGNORECASE)
                for segmento in segmentos:
                    texto = self._limpar_html(segmento)
                    if texto:
                        elementos.append({"tipo": "texto", "conteudo": texto})

        return elementos

    def _criar_tabela_xml(self, linhas: list):
        """Cria um elemento w:tbl XML com bordas a partir dos dados das linhas."""
        num_cols = max(len(row) for row in linhas) if linhas else 1

        tbl = OxmlElement('w:tbl')

        # Propriedades da tabela
        tblPr = OxmlElement('w:tblPr')

        tblStyle = OxmlElement('w:tblStyle')
        tblStyle.set(qn('w:val'), 'TableGrid')
        tblPr.append(tblStyle)

        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        tblPr.append(tblW)

        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        tblPr.append(jc)

        # Bordas
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        tbl.append(tblPr)

        # Grid de colunas
        tblGrid = OxmlElement('w:tblGrid')
        col_width = str(9000 // num_cols)
        for _ in range(num_cols):
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), col_width)
            tblGrid.append(gridCol)
        tbl.append(tblGrid)

        # Linhas e células
        for row_data in linhas:
            tr = OxmlElement('w:tr')
            for i in range(num_cols):
                tc = OxmlElement('w:tc')
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = row_data[i] if i < len(row_data) else ''
                t.set(qn('xml:space'), 'preserve')
                r.append(t)
                p.append(r)
                tc.append(p)
                tr.append(tc)
            tbl.append(tr)

        return tbl

    def _inserir_elementos_no_documento(self, doc_path: str, subcategorias_elementos: list):
        """
        Insere elementos (textos e tabelas) das subcategorias no documento,
        logo após o parágrafo onde o ${block_name} foi substituído.

        subcategorias_elementos: lista de listas de elementos, ex:
        [
            [{"tipo": "texto", ...}, {"tipo": "tabela", ...}],  # subcategoria 1
            [{"tipo": "texto", ...}],                            # subcategoria 2
        ]
        """
        if not subcategorias_elementos:
            return

        doc = Document(doc_path)

        # Localiza o parágrafo da primeira subcategoria (texto do ${block_name})
        primeiro_texto = ""
        for elem in subcategorias_elementos[0]:
            if elem["tipo"] == "texto":
                primeiro_texto = elem["conteudo"]
                break

        if not primeiro_texto:
            del doc
            gc.collect()
            return

        para_ref = None
        for para in doc.paragraphs:
            if primeiro_texto[:80] in para.text:
                para_ref = para
                break

        if not para_ref:
            del doc
            gc.collect()
            return

        # Primeira subcategoria: pula o primeiro texto (já está no ${block_name})
        todos_elementos = []
        first_text_skipped = False
        for elem in subcategorias_elementos[0]:
            if not first_text_skipped and elem["tipo"] == "texto":
                first_text_skipped = True
                continue
            todos_elementos.append(elem)

        # Subcategorias restantes: todos os elementos
        for subcat_elems in subcategorias_elementos[1:]:
            todos_elementos.extend(subcat_elems)

        if not todos_elementos:
            del doc
            gc.collect()
            return

        # Insere sequencialmente, rastreando o ponto de inserção
        insert_after = para_ref._element
        for elem in todos_elementos:
            if elem["tipo"] == "texto":
                new_p = deepcopy(para_ref._element)
                runs = new_p.findall(qn('w:r'))
                if runs:
                    for r in runs[1:]:
                        new_p.remove(r)
                    ts = runs[0].findall(qn('w:t'))
                    if ts:
                        ts[0].text = elem["conteudo"]
                        ts[0].set(qn('xml:space'), 'preserve')
                insert_after.addnext(new_p)
                insert_after = new_p

            elif elem["tipo"] == "tabela":
                tbl = self._criar_tabela_xml(elem["linhas"])
                insert_after.addnext(tbl)
                insert_after = tbl

        doc.save(doc_path)
        del doc
        gc.collect()

    # ──────────────────────────────────────────────────────────────────────
    # Fluxo principal: criar documento
    # ──────────────────────────────────────────────────────────────────────
    def criar_documento(self, tipo_documento: str, dados: Dict[str, str],
                        subcategorias_texto: List[str] = None,
                        output_path: str = None) -> str:
        """
        Fluxo de criação de documento:
        1. Baixa o .docx do link S3 (tipo_documento = URL direta)
        2. Monta o dicionário ${VAR}: valor a partir dos dados já prontos
        3. Limpa HTML dos textos de subcategorias e mapeia para ${block_name}
        4. Substitui variáveis no template
        5. Insere subcategorias adicionais como parágrafos separados

        Args:
            tipo_documento: URL direta do template .docx no S3
            dados: Dicionário {NOME_VARIAVEL: valor} já pronto
                   Ex: {"NOME_DO_REU": "João", "VARA": "1"}
            subcategorias_texto: Lista de textos HTML das subcategorias,
                                 já na ordem correta.
            output_path: Caminho de saída. Se None, gera automaticamente.

        Returns:
            Caminho do arquivo .docx gerado.
        """
        # 1. Baixar template direto da URL S3
        template_path = self.baixar_template(tipo_documento)

        try:
            # 2. Montar dicionário ${VAR}: valor
            dados_completos = {}
            for var_nome, valor in dados.items():
                if valor == "GETDATE()":
                    dados_completos[f"${{{var_nome}}}"] = self.db.select_data_atual()
                elif var_nome == "VARA":
                    dados_completos[f"${{{var_nome}}}"] = f"{valor}ª"
                else:
                    dados_completos[f"${{{var_nome}}}"] = valor

            # Remover variáveis ignoradas que existam no template
            for var_ignorada in self.VARIAVEIS_IGNORADAS:
                dados_completos[f"${{{var_ignorada}}}"] = ""

            # 3. Processar subcategorias (texto + tabelas)
            subcategorias_elementos = []
            if subcategorias_texto:
                for texto_html in subcategorias_texto:
                    subcategorias_elementos.append(self._processar_html(texto_html))

                # Para ${block_name}, usa apenas o primeiro texto da primeira subcategoria
                primeiro_texto = ""
                for elem in subcategorias_elementos[0]:
                    if elem["tipo"] == "texto":
                        primeiro_texto = elem["conteudo"]
                        break

                dados_completos["${block_name}"] = primeiro_texto
                dados_completos["${/block_name}"] = ""

            # 4. Substituir variáveis no template
            if output_path is None:
                output_path = "documento_GERADO.docx"

            total = self.substituir_variaveis(template_path, dados_completos, output_path)
            logger.info(f"[GerarPeticao] {total} substituicoes realizadas")

            # 5. Inserir elementos restantes (textos e tabelas) das subcategorias
            if subcategorias_elementos:
                tem_extras = (
                    len(subcategorias_elementos) > 1
                    or any(e["tipo"] == "tabela" for e in subcategorias_elementos[0])
                    or sum(1 for e in subcategorias_elementos[0] if e["tipo"] == "texto") > 1
                )
                if tem_extras:
                    self._inserir_elementos_no_documento(output_path, subcategorias_elementos)
        finally:
            # Sempre apaga o template baixado, mesmo se der erro
            if os.path.exists(template_path):
                try:
                    os.unlink(template_path)
                except PermissionError:
                    gc.collect()
                    os.unlink(template_path)

        return output_path

